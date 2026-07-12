from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "application" / "BidRadar-X_开题补充材料_初稿.docx"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 103, 115)
LIGHT = "F2F4F7"
ACCENT = "E8EEF5"
INK = RGBColor(24, 33, 43)


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_first_row_as_header(table):
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def style_table(table, widths_dxa, header_fill=LIGHT):
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    mark_first_row_as_header(table)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_font(run, 9.5, bold=True, color=DARK)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_font(run, 9)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    mark_first_row_as_header(table)
    set_cell_shading(table.cell(0, 0), ACCENT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(f"{label}  "), 10.5, bold=True, color=DARK)
    set_font(p.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("BidRadar-X  |  开题补充材料"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("AI先锋未来人才大赛 · 超聚变命题 · 初稿"), 8.5, color=MUTED)

    # Page 1: proposal centerpiece
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(56)
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("超聚变企业命题开题补充材料"), 12, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("BidRadar-X"), 28, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_font(p.add_run("面向服务器与算力场景的可信招投标机会雷达"), 15, color=MUTED)
    add_callout(doc, "一句话方案", "自然语言创建立即或定时任务，聚合公开公告，识别同一项目与版本变化，并生成字段可回查的Word报告。")
    add_heading(doc, "Executive Summary", 1)
    add_bullet(doc, "问题不是公告不足，而是信息分散、重复、持续变化，难以快速形成可信的项目判断。")
    add_bullet(doc, "聚合与关键词订阅已是成熟能力；项目以垂直语义、跨站归并、变化优先和字段证据链形成差异。")
    add_bullet(doc, "创新通过抓取成功率、Precision@10、去重F1和证据覆盖率验证，不以Agent数量证明效果。")
    add_body(doc, "材料状态：报名阶段方案初稿。本文所有系统效果数字均为后续试点验收目标，并非已取得成绩。")

    # Page 2
    page_break(doc)
    add_heading(doc, "1. 为什么值得解决", 1)
    add_body(doc, "市场规模大且信息持续更新。财政部国库司披露，2024年全国政府采购规模为33,750.43亿元，其中公开招标占76.63%。中国政府采购网每天持续发布中央和地方的招标、更正、中标等公告。", "市场规模大且信息持续更新。")
    add_body(doc, "真正瓶颈在项目判断。业务人员不仅要找到公告，还要判断它是否与服务器、液冷或算力中心相关，是否与其他站点记录重复，截止时间或金额是否发生变化，以及报告中的结论能否回到原文核验。", "真正瓶颈在项目判断。")
    add_heading(doc, "竞品说明：聚合和订阅只是入场券", 2)
    table = doc.add_table(rows=1, cols=4)
    for i, value in enumerate(("能力", "剑鱼/千里马公开能力", "BidRadar-X重点", "判断")):
        table.cell(0, i).text = value
    rows = [
        ("搜索与订阅", "关键词、地区、类型、定时推送", "自然语言转结构化任务", "基础能力"),
        ("分析", "业主、企业、竞争对手、市场分析", "算力行业相关性与项目时间线", "垂直聚焦"),
        ("可信度", "提供原文链接和附件", "每个关键字段绑定原文片段与哈希", "核心差异"),
        ("变化追踪", "项目/企业监控", "跨站归并后只报告新增与变化", "核心差异"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, [1500, 3000, 3000, 1860])
    p = doc.add_paragraph("注：竞品判断仅依据公开产品页面；未公开展示的能力不能据此断言不存在。")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(""), 9, italic=True, color=MUTED)
    add_callout(doc, "产品结论", "从“搜得多”升级为“找得准、看得懂、可核验、能追踪”。")

    # Page 3
    page_break(doc)
    add_heading(doc, "2. 从一句需求到一份可核验报告", 1)
    add_number(doc, "解析任务：将地区、主题、公告类型、时间、频率和输出要求转为JobSpec；歧义条件要求用户确认。")
    add_number(doc, "决定执行：无时间要求立即执行；明确频率交给调度器，保存的是结构化任务而非原始提示词。")
    add_number(doc, "合规采集：SourceAdapter接入无需绕过登录或验证码的公开站点，保存网页、PDF、抓取时间和内容哈希。")
    add_number(doc, "结构化处理：解析字段、进行行业相关性排序，并把跨站、跨阶段公告归入同一项目。")
    add_number(doc, "证据验证：预算、截止时间、采购人等字段必须绑定原文片段；冲突与缺失显式标记。")
    add_number(doc, "输出与追踪：生成Word和任务记录，定时报告优先展示新增、延期、更正、废标和中标。")
    add_heading(doc, "模型只做适合模型的工作", 2)
    add_bullet(doc, "允许：意图解析、非关键文本字段的辅助提取、基于已验证记录的摘要。")
    add_bullet(doc, "禁止依赖模型：网络访问、限速重试、任务状态、项目版本、证据核验和Word表格生成。")
    add_body(doc, "这样的边界使系统在模型超时或输出非法JSON时仍可降级，并让每个关键结果都有可重复测试。")

    # Page 4
    page_break(doc)
    add_heading(doc, "3. 两项创新直接对应业务风险", 1)
    add_heading(doc, "创新一｜字段级证据链", 2)
    add_body(doc, "每个关键字段保存字段值、原文片段、来源URL、抓取时间、内容哈希和验证状态。字段缺失时统一显示“未知”，模型不能根据上下文补写。评委可以从报告字段回到原公告核验。")
    add_callout(doc, "解决的风险", "降低摘要看似完整、但预算或截止日期无法证明的风险。")
    add_heading(doc, "创新二｜项目身份与变更雷达", 2)
    add_body(doc, "系统以项目编号、标题、采购人、地区和时间生成项目候选对，把预告、招标、更正、废标和中标公告组织为版本时间线。定时报告不重复堆叠旧公告，而是优先展示新增、金额变化、截止时间变化和状态变化。")
    add_callout(doc, "解决的风险", "减少跨站重复阅读和关键变更被后续公告淹没的风险。")
    add_heading(doc, "演示场景", 2)
    add_body(doc, "“每天8:30，搜索上海、江苏、浙江近24小时内与服务器、液冷、算力中心有关的公告，排除办公家具，只报告高相关或发生变化的项目并生成Word。”")

    # Page 5
    page_break(doc)
    add_heading(doc, "4. 用指标证明，不用概念包装", 1)
    table = doc.add_table(rows=1, cols=4)
    for i, value in enumerate(("指标", "定义", "目标", "基线/样本")):
        table.cell(0, i).text = value
    metric_rows = [
        ("抓取成功率", "成功获得可解析详情的公告/目标公告", "≥90%", "两轮×30条/来源"),
        ("Precision@10", "前10条中人工判定相关的比例", "≥80%", "关键词排序基线"),
        ("去重配对F1", "同项目候选对的精确率与召回率调和", "≥0.85", "人工重复对"),
        ("证据覆盖率", "报告关键字段有可定位证据的比例", "100%", "全部报告字段"),
        ("任务耗时", "陌生用户从输入到报告下载", "≤15分钟", "可用性走查"),
    ]
    for row in metric_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, [1800, 3600, 1260, 2700])
    add_heading(doc, "落地可行性", 2)
    add_bullet(doc, "技术：FastAPI、APScheduler、SQLite、本地Web和DOCX；两人可按模块并行。")
    add_bullet(doc, "数据：优先验证中国政府采购网、全国公共资源交易平台、上海政府采购网；未通过合规门槛的不进入MVP。")
    add_bullet(doc, "演示：实时源不可用时切换到相同契约的固定样例，确保测试和现场流程可复现。")
    add_bullet(doc, "推广：SourceAdapter和行业词库可配置，可从华东算力场景扩展到医疗设备、能源或工程服务。")

    # Page 6
    page_break(doc)
    add_heading(doc, "5. 团队、计划与证据补齐", 1)
    add_heading(doc, "两人分工", 2)
    add_bullet(doc, "数据科学/产品负责人：行业词库、字段标准、相关性、去重、标注评测、证据规则和业务材料。")
    add_bullet(doc, "软件工程/平台负责人：采集适配器、解析、任务状态、调度、存储、API、Web和Word生成。")
    add_bullet(doc, "共同负责：公共契约、样本复核、端到端验收、演示与风险说明。")
    add_heading(doc, "报名前优先补齐", 2)
    add_number(doc, "访谈3–5名销售运营、采购信息检索或投标支持相关人员，匿名记录场景、频率和痛点。")
    add_number(doc, "验证三个候选来源的条款、robots、验证码、限速和30条小样，不绕过任何访问控制。")
    add_number(doc, "制作10–20条端到端固定样例，生成第一份Word并记录失败路径。")
    add_heading(doc, "来源与边界", 2)
    sources = [
        "财政部国库司：《2024年全国政府采购简要情况》，中国政府采购网。",
        "中国政府采购网：采购公告公开页面。",
        "剑鱼标讯：免费订阅及产品服务公开页面。",
        "千里马招标网：产品功能公开介绍页面。",
    ]
    for source in sources:
        add_bullet(doc, source)
    add_body(doc, "重要边界：当前尚未完成真实从业者访谈，也尚未取得系统评测结果。所有目标值将在人工核验集上验证；登录平台未经授权不接入。")

    doc.core_properties.title = "BidRadar-X 开题补充材料"
    doc.core_properties.subject = "AI先锋未来人才大赛超聚变命题报名材料"
    doc.core_properties.author = "BidRadar-X Team"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
