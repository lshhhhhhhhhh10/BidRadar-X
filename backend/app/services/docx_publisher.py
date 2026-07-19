"""Generate compact, evidence-linked Word reports for tender notices."""

from __future__ import annotations

from collections.abc import Callable, Sequence
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
import re
from typing import Any, Literal
from zoneinfo import ZoneInfo

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from ..schemas.tender import EvidenceReference, TenderNotice


ReportScope = Literal["full", "incremental"]

_SCOPE_LABELS: dict[ReportScope, str] = {
    "full": "全量报告",
    "incremental": "仅新增内容报告",
}
_SECTION_ORDER = {
    "procurement": 0,
    "qualification": 1,
    "technical": 2,
    "timeline": 3,
    "commercial": 4,
    "submission": 5,
    "evaluation": 6,
    "reference": 7,
}
_WINDOWS_INVALID = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')
_WINDOWS_RESERVED = {
    "CON", "PRN", "AUX", "NUL",
    *(f"COM{number}" for number in range(1, 10)),
    *(f"LPT{number}" for number in range(1, 10)),
}
_CJK_FONT = "Arial Unicode MS"
_TABLE_WIDTH_DXA = 9360
_TABLE_INDENT_DXA = 120
_RISK_LABELS = {
    "low": ("低", "2E6B4F", "EAF4EE"),
    "medium": ("中", "7A5A00", "FFF6DB"),
    "high": ("高", "9B1C1C", "FCEBEC"),
    "unknown": ("待核验", "52606D", "EEF1F4"),
}


@dataclass(frozen=True)
class DocxValidationResult:
    """Observable structure recovered by reopening a generated DOCX."""

    paragraph_count: int
    table_count: int
    hyperlink_targets: tuple[str, ...]


class DocxValidationError(ValueError):
    """Raised when a generated DOCX cannot be reopened or fails its contract."""


class DocxPublisher:
    """Render a standalone Word report without invoking the workflow."""

    def __init__(
        self,
        output_dir: Path,
        clock: Callable[[], datetime] | None = None,
    ) -> None:
        self.output_dir = Path(output_dir)
        self._clock = clock or (lambda: datetime.now(ZoneInfo("Asia/Shanghai")))

    def publish(
        self,
        query: str,
        notices: Sequence[TenderNotice],
        report_scope: ReportScope = "full",
        ai_report: dict[str, Any] | None = None,
    ) -> Path:
        """Generate, exclusively write, reopen, and validate one DOCX report."""

        if report_scope not in _SCOPE_LABELS:
            raise ValueError("report_scope must be 'full' or 'incremental'")
        notice_list = list(notices)
        if any(not isinstance(notice, TenderNotice) for notice in notice_list):
            raise TypeError("notices must contain only TenderNotice instances")

        generated_at = self._clock()
        filename = build_report_filename(query, generated_at)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        report_path = self.output_dir / filename
        document = _build_document(
            query,
            notice_list,
            report_scope,
            generated_at,
            ai_report=ai_report,
        )
        buffer = BytesIO()
        document.save(buffer)

        try:
            report_file = report_path.open("xb")
        except FileExistsError as error:
            raise FileExistsError(
                f"report already exists for this query and minute: {report_path.name}"
            ) from error
        try:
            with report_file:
                report_file.write(buffer.getvalue())
        except Exception:
            report_path.unlink(missing_ok=True)
            raise

        try:
            validate_docx(
                report_path,
                expected_notices=notice_list,
                expected_scope=report_scope,
            )
        except Exception:
            report_path.unlink(missing_ok=True)
            raise
        return report_path


def build_report_filename(
    query: str,
    generated_at: datetime,
    *,
    project_sequence: int | None = None,
) -> str:
    """Build {question}_{YYYYMMDDHHmm}.docx with an optional project suffix."""

    safe_query = _WINDOWS_INVALID.sub("_", query)
    safe_query = re.sub(r"\s+", " ", safe_query).strip(" .")
    safe_query = re.sub(r"_+", "_", safe_query).strip(" ._")
    if not safe_query:
        safe_query = "情报任务"
    if safe_query.split(".", maxsplit=1)[0].upper() in _WINDOWS_RESERVED:
        safe_query = f"_{safe_query}"
    safe_query = safe_query[:80].rstrip(" ._") or "情报任务"
    sequence_suffix = (
        f"_项目{project_sequence:02d}"
        if project_sequence is not None and project_sequence > 0
        else ""
    )
    return f"{safe_query}_{generated_at.strftime('%Y%m%d%H%M')}{sequence_suffix}.docx"


def validate_docx(
    report_path: Path,
    *,
    expected_notices: Sequence[TenderNotice] | None = None,
    expected_scope: ReportScope | None = None,
) -> DocxValidationResult:
    """Reopen a DOCX and validate its core structure and external links."""

    try:
        document = Document(report_path)
    except Exception as error:
        raise DocxValidationError(f"cannot reopen generated DOCX: {report_path}") from error

    paragraphs = list(document.paragraphs)
    paragraphs.extend(
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    if not paragraphs or not document.tables:
        raise DocxValidationError("generated DOCX is missing paragraphs or tables")

    hyperlink_targets: list[str] = []
    for hyperlink in document.element.xpath(".//w:hyperlink"):
        relationship_id = hyperlink.get(qn("r:id"))
        relationship = document.part.rels.get(relationship_id)
        if relationship is None or relationship.reltype != RT.HYPERLINK:
            raise DocxValidationError("DOCX contains an invalid hyperlink relationship")
        target = relationship.target_ref
        if not relationship.is_external or not target.startswith(("http://", "https://")):
            raise DocxValidationError(f"DOCX hyperlink is not HTTP(S): {target}")
        hyperlink_targets.append(target)

    all_text = "\n".join(paragraph.text for paragraph in paragraphs)
    if expected_scope is not None and _SCOPE_LABELS[expected_scope] not in all_text:
        raise DocxValidationError("generated DOCX is missing the requested report scope")
    if expected_notices is not None:
        remaining_tables = _notice_tables(document)
        for notice in expected_notices:
            if notice.core_content not in all_text:
                raise DocxValidationError(
                    f"generated DOCX is missing required core content: {notice.title}"
                )
            match = next(
                (
                    index for index, table in enumerate(remaining_tables)
                    if _notice_table_matches(document, table, notice)
                ),
                None,
            )
            if match is None:
                raise DocxValidationError(
                    f"generated DOCX is missing the required notice table: {notice.title}"
                )
            remaining_tables.pop(match)

    return DocxValidationResult(
        paragraph_count=len(paragraphs),
        table_count=len(document.tables),
        hyperlink_targets=tuple(hyperlink_targets),
    )


def _build_document(
    query: str,
    notices: list[TenderNotice],
    report_scope: ReportScope,
    generated_at: datetime,
    ai_report: dict[str, Any] | None = None,
) -> DocumentType:
    document = Document()
    _configure_document(document, generated_at)
    _add_title_page(document, query, notices, report_scope, generated_at)

    if not notices:
        document.add_paragraph("未发现符合条件的公告。")
        return document

    ordered = sorted(notices, key=lambda item: (item.published_at, item.title), reverse=True)
    _add_overview(document, ordered)
    if ai_report and ai_report.get("status") == "generated":
        _add_ai_overview(document, ai_report)
    else:
        _add_ai_unavailable_note(document, ai_report)
    narratives = {
        item["notice_id"]: item
        for item in (ai_report or {}).get("notice_narratives", [])
        if isinstance(item, dict) and item.get("notice_id")
    }
    for index, notice in enumerate(ordered, start=1):
        if index > 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _add_project(document, index, notice, narratives.get(notice.notice_id))
    return document


def _add_ai_overview(document: DocumentType, ai_report: dict[str, Any]) -> None:
    document.add_heading("AI 辅助摘要与风险研判", level=1)
    _add_callout(
        document,
        "AI 辅助摘要",
        str(ai_report.get("executive_summary", "")),
        fill="F4F6F9",
        accent="1F4D78",
    )

    narratives = [
        item for item in ai_report.get("notice_narratives", [])
        if isinstance(item, dict)
    ]
    risk_counts = {key: 0 for key in _RISK_LABELS}
    for item in narratives:
        level = str(item.get("risk_level") or "unknown")
        risk_counts[level if level in risk_counts else "unknown"] += 1
    if narratives:
        risk_line = document.add_paragraph()
        risk_line.add_run("风险分布：").bold = True
        risk_line.add_run(
            "  ".join(
                f"{_RISK_LABELS[key][0]} {risk_counts[key]} 项"
                for key in ("high", "medium", "low", "unknown")
                if risk_counts[key]
            )
        )
        _style_paragraph(risk_line, 10, "263238", spacing_before=4, spacing_after=8)

    findings = ai_report.get("key_findings", [])
    if findings:
        document.add_heading("跨公告关键发现", level=2)
        for item in findings:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(str(item.get("text", "")))
            evidence_ids = item.get("evidence_ids", [])
            if evidence_ids:
                paragraph.add_run(f"  [证据：{', '.join(evidence_ids)}]")
            _style_paragraph(paragraph, 10, "263238", spacing_after=4)

    note = document.add_paragraph(
        "说明：以上内容由模型基于已采集证据生成，仅用于辅助筛选；项目事实、原文链接与附件链接以公告原文为准。"
    )
    _style_paragraph(note, 9, "6B7C8A", spacing_before=4, spacing_after=8)


def _add_ai_unavailable_note(
    document: DocumentType,
    ai_report: dict[str, Any] | None,
) -> None:
    document.add_heading("AI 辅助摘要与风险研判", level=1)
    reason = str((ai_report or {}).get("reason") or "AI 接口未启用或输出未通过校验")
    _add_callout(
        document,
        "本次未生成 AI 辅助分析",
        f"{reason}。为避免把未经模型校验的内容伪装成 AI 结论，本报告仅保留原公告事实与证据。",
        fill="FFF6DB",
        accent="7A5A00",
    )


def _configure_document(document: DocumentType, generated_at: datetime) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    _style_font(normal, size=11, color="263238")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in (
        ("Title", 24, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = document.styles[style_name]
        _style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.keep_with_next = True
        before_after = {
            "Title": (0, 8),
            "Heading 1": (16, 8),
            "Heading 2": (12, 6),
            "Heading 3": (8, 4),
        }[style_name]
        style.paragraph_format.space_before = Pt(before_after[0])
        style.paragraph_format.space_after = Pt(before_after[1])

    for list_style_name in ("List Bullet", "List Number"):
        list_style = document.styles[list_style_name]
        _style_font(list_style, size=11, color="263238")
        list_style.paragraph_format.left_indent = Inches(0.5)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)
        list_style.paragraph_format.space_after = Pt(8)
        list_style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = (
        "BIDRADAR X  ·  招投标机会研判"
        f"                                      {generated_at.strftime('%Y.%m.%d')}"
    )
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _style_paragraph(header, 8.5, "6B7C8A", bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("BidRadar-X  |  ")
    _add_page_number(footer)
    _style_paragraph(footer, 8, "7D8993")


def _add_title_page(
    document: DocumentType,
    query: str,
    notices: list[TenderNotice],
    report_scope: ReportScope,
    generated_at: datetime,
) -> None:
    kicker = document.add_paragraph("招投标机会研判  ·  证据可追溯")
    _style_paragraph(kicker, 9, "2E74B5", bold=True, spacing_after=6)
    title = document.add_paragraph(style="Title")
    title.add_run("项目机会分析报告")
    subtitle = document.add_paragraph(query)
    _style_paragraph(subtitle, 15, "172B4D", bold=True, spacing_after=14)

    metadata = document.add_table(rows=1, cols=4)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_dxa = (2340, 2340, 2340, 2340)
    _set_table_geometry(metadata, widths_dxa)
    metrics = (
        ("报告范围", _SCOPE_LABELS[report_scope]),
        ("项目数量", f"{len(notices)} 项"),
        ("生成时间", generated_at.strftime("%Y-%m-%d %H:%M")),
        ("事实口径", "公告原文与附件"),
    )
    for cell, (label, value) in zip(metadata.rows[0].cells, metrics, strict=True):
        _shade_cell(cell, "F2F4F7")
        _set_cell_margins(cell, 120, 140, 120, 140)
        label_paragraph = _clear_cell(cell)
        label_paragraph.add_run(label)
        _style_paragraph(label_paragraph, 8.5, "6B7C8A", spacing_after=3)
        value_paragraph = cell.add_paragraph(value)
        _style_paragraph(value_paragraph, 10, "172B4D", bold=True)

    _add_callout(
        document,
        "报告口径",
        (
            "原公告事实字段由程序锁定，不交由模型改写；AI 仅生成辅助摘要、机会提示和风险研判，并绑定证据编号。点击原公告或附件链接可回溯核验。"
            if notices
            else "本次没有符合条件的公告，因此不生成项目事实、AI 研判或附件入口。"
        ),
        fill="F4F6F9",
        accent="1F4D78",
    )


def _add_overview(document: DocumentType, notices: list[TenderNotice]) -> None:
    document.add_heading("一页速览", level=1)
    nearest_deadline = min(
        (notice.deadline for notice in notices if notice.deadline is not None),
        default=None,
    )
    known_budget_notices = [notice for notice in notices if notice.budget is not None]
    budget_currencies = {notice.budget_currency for notice in known_budget_notices}
    if known_budget_notices and len(budget_currencies) == 1:
        currency = next(iter(budget_currencies))
        budget_summary = f"{sum(notice.budget or 0 for notice in known_budget_notices):,.0f} {currency}"
    elif known_budget_notices:
        budget_summary = "多币种，见项目事实"
    else:
        budget_summary = "原文未载明"
    source_count = len({notice.source.source_name for notice in notices})
    metrics = (
        ("公告数量", f"{len(notices)} 条"),
        ("最近截止", nearest_deadline.strftime("%m-%d %H:%M") if nearest_deadline else "原文未载明"),
        ("已披露预算", budget_summary),
        ("证据来源", f"{source_count} 个"),
    )
    metric_table = document.add_table(rows=1, cols=4)
    metric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metric_widths = (2340, 2340, 2340, 2340)
    _set_table_geometry(metric_table, metric_widths)
    for cell, (label, value) in zip(metric_table.rows[0].cells, metrics, strict=True):
        _shade_cell(cell, "F2F4F7")
        _set_cell_margins(cell, 120, 140, 120, 140)
        label_paragraph = _clear_cell(cell)
        label_paragraph.add_run(label)
        _style_paragraph(label_paragraph, 8.5, "6B7C8A", spacing_after=3)
        value_paragraph = cell.add_paragraph(value)
        _style_paragraph(value_paragraph, 11, "172B4D", bold=True)

    document.add_heading("公告清单", level=2)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_dxa = (5280, 2160, 1920)
    _set_table_geometry(table, widths_dxa)
    for cell, text in zip(
        table.rows[0].cells,
        ("项目名称", "关键时间", "信息来源"),
        strict=True,
    ):
        cell.text = text
        _shade_cell(cell, "E8EFF4")
        _style_cell(cell, 8.5, "1F4D78", bold=True)
    for index, notice in enumerate(notices, start=1):
        row = table.add_row()
        values = (
            f"{index}. {notice.title}",
            notice.deadline.strftime("截止 %m-%d %H:%M") if notice.deadline else f"发布 {notice.published_at.strftime('%m-%d')}",
            notice.source.source_name,
        )
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
            _set_cell_margins(cell, 80, 120, 80, 120)
            _style_cell(cell, 9, "263238")
        if index % 2 == 0:
            for cell in row.cells:
                _shade_cell(cell, "F4F7F9")
    _set_table_geometry(table, widths_dxa)


def _add_project(
    document: DocumentType,
    index: int,
    notice: TenderNotice,
    ai_narrative: dict[str, Any] | None = None,
) -> None:
    marker = document.add_paragraph(f"项目 {index:02d}  ·  原公告事实")
    _style_paragraph(marker, 8.5, "4C7899", bold=True, spacing_after=2)
    document.add_heading(notice.title, level=1)
    context = document.add_paragraph(
        f"{notice.source.source_name}  ·  {notice.published_at.strftime('%Y-%m-%d')}"
        f"  ·  {notice.region or '区域未载明'}"
    )
    _style_paragraph(context, 9, "6B7C8A", spacing_after=8)

    document.add_heading("项目事实", level=2)
    _add_notice_summary(document, notice)

    if ai_narrative:
        document.add_heading("AI 辅助研判", level=2)
        _add_callout(
            document,
            "AI 辅助摘要",
            str(ai_narrative.get("summary", "")),
            fill="F4F6F9",
            accent="1F4D78",
        )
        risk_level = str(ai_narrative.get("risk_level") or "unknown")
        risk_label, risk_color, risk_fill = _RISK_LABELS.get(
            risk_level, _RISK_LABELS["unknown"]
        )
        risk_table = document.add_table(rows=1, cols=2)
        _set_table_geometry(risk_table, (1800, 7560))
        risk_table.cell(0, 0).text = f"风险等级\n{risk_label}"
        risk_table.cell(0, 1).text = str(
            ai_narrative.get("risk_assessment") or "证据不足，需人工复核。"
        )
        _shade_cell(risk_table.cell(0, 0), risk_fill)
        _shade_cell(risk_table.cell(0, 1), "FFFFFF")
        _style_cell(risk_table.cell(0, 0), 10, risk_color, bold=True)
        _style_cell(risk_table.cell(0, 1), 10, "263238")
        for cell in risk_table.rows[0].cells:
            _set_cell_margins(cell, 100, 120, 100, 120)

        _add_bullet_group(document, "风险提示", ai_narrative.get("risk_points", []))
        _add_bullet_group(document, "机会提示", ai_narrative.get("opportunity_points", []))
        _add_bullet_group(document, "建议动作", ai_narrative.get("next_actions", []))
        evidence_ids = ai_narrative.get("evidence_ids", [])
        if evidence_ids:
            evidence_line = document.add_paragraph(f"关联证据：{', '.join(evidence_ids)}")
            _style_paragraph(evidence_line, 8.5, "6B7C8A", spacing_after=8)

    document.add_heading("原公告核心内容", level=2)
    summary = document.add_paragraph(notice.core_content)
    _style_paragraph(summary, 10.5, "263238", spacing_after=8)

    sections = sorted(
        notice.requirement_sections,
        key=lambda section: _SECTION_ORDER.get(section.section_id, 99),
    )
    if sections:
        document.add_heading("招标条款与证据", level=1)
    for section_index, section in enumerate(sections, start=1):
        _add_requirement_section(document, section_index, section, notice.evidence)


def _add_notice_summary(document: DocumentType, notice: TenderNotice) -> None:
    table = document.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_geometry(table, (1800, 2880, 1800, 2880))
    pairs = (
        ("项目名称", notice.title, "公告类型", _notice_kind_label(notice)),
        ("招标人", notice.purchaser or "原公告未载明", "项目编号", notice.project_code or "原公告未载明"),
        ("发布时间", notice.published_at.strftime("%Y-%m-%d %H:%M"), "投标截止", notice.deadline.strftime("%Y-%m-%d %H:%M") if notice.deadline else "原公告未载明"),
        ("项目区域", notice.region or "原公告未载明", "预算金额", _budget_label(notice)),
        ("信息来源", notice.source.source_name, "原公告", None),
        ("附件", None, "采集时间", notice.fetched_at.strftime("%Y-%m-%d %H:%M")),
    )
    for row, values in zip(table.rows, pairs, strict=True):
        for cell in row.cells:
            _set_cell_margins(cell, 90, 120, 90, 120)
        for label_index in (0, 2):
            row.cells[label_index].text = values[label_index]
            _shade_cell(row.cells[label_index], "E8EFF4")
            _style_cell(row.cells[label_index], 9, "1F4D78", bold=True)
        for value_index in (1, 3):
            value = values[value_index]
            if value is not None:
                row.cells[value_index].text = value
                _style_cell(row.cells[value_index], 9.5, "263238")

    source_paragraph = _clear_cell(table.cell(4, 3))
    _add_hyperlink(source_paragraph, str(notice.source.source_url), "查看原公告 ↗")
    attachment_paragraph = _clear_cell(table.cell(5, 1))
    if notice.attachments:
        for index, attachment in enumerate(notice.attachments):
            if index:
                attachment_paragraph.add_run("；")
            _add_hyperlink(
                attachment_paragraph,
                str(attachment.url),
                attachment.name or "查看附件",
            )
    else:
        attachment_paragraph.add_run("无")


def _add_requirement_section(document, index, section, evidence_items) -> None:
    title = section.title
    if section.section_id == "technical" and "技术要求" not in title:
        title = f"{title}（技术要求）"
    document.add_heading(f"{index}. {title}", level=2)
    if section.summary:
        summary = document.add_paragraph(section.summary)
        _style_paragraph(summary, 8.5, "6B7C8A", spacing_after=4)

    evidence_by_id = {item.evidence_id: item for item in evidence_items}
    known_facts = [fact for fact in section.facts if fact.value is not None]
    if not known_facts and not section.tables:
        missing = document.add_paragraph("本公告未提及可核验条款。")
        _style_paragraph(missing, 8.5, "88959E")
        return

    if known_facts:
        table = document.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths_dxa = (1600, 3080, 1600, 3080)
        _set_table_geometry(table, widths_dxa)
        for cell, label in zip(
            table.rows[0].cells,
            ("要点", "识别原文", "要点", "识别原文"),
            strict=True,
        ):
            cell.text = label
            _shade_cell(cell, "DCE7EE")
            _style_cell(cell, 8, "315D7D", bold=True)
        for offset in range(0, len(known_facts), 2):
            row = table.add_row()
            for cell in row.cells:
                _set_cell_margins(cell, 80, 120, 80, 120)
            for pair_index, fact in enumerate(known_facts[offset:offset + 2]):
                label_cell = row.cells[pair_index * 2]
                value_cell = row.cells[pair_index * 2 + 1]
                label_cell.text = fact.label
                _shade_cell(label_cell, "F2F6F8")
                _style_cell(label_cell, 9, "1F4D78", bold=True)
                value_cell.text = fact.value or ""
                _style_cell(value_cell, 9, "263238")
                _append_fact_source(value_cell.paragraphs[0], fact.evidence_ids, evidence_by_id)
        _set_table_geometry(table, widths_dxa)

    for requirement_table in section.tables:
        document.add_heading(requirement_table.title, level=3)
        table = document.add_table(rows=1, cols=len(requirement_table.columns) + 1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        column_count = len(requirement_table.columns) + 1
        base_width = _TABLE_WIDTH_DXA // column_count
        widths_dxa = [base_width] * column_count
        widths_dxa[-1] += _TABLE_WIDTH_DXA - sum(widths_dxa)
        _set_table_geometry(table, tuple(widths_dxa))
        for cell, label in zip(
            table.rows[0].cells,
            [*requirement_table.columns, "证据"],
            strict=True,
        ):
            cell.text = label
            _shade_cell(cell, "DCE7EE")
            _style_cell(cell, 8, "315D7D", bold=True)
        for requirement_row in requirement_table.rows:
            cells = table.add_row().cells
            for cell, value in zip(cells[:-1], requirement_row.cells, strict=True):
                cell.text = value or "原公告未载明"
                _style_cell(cell, 9, "263238")
            _append_fact_source(
                cells[-1].paragraphs[0],
                requirement_row.evidence_ids,
                evidence_by_id,
            )
        _set_table_geometry(table, tuple(widths_dxa))


def _append_fact_source(paragraph, evidence_ids, evidence_by_id) -> None:
    matching = [evidence_by_id[item] for item in evidence_ids if item in evidence_by_id]
    if not matching:
        paragraph.add_run("无")
        return
    for index, evidence in enumerate(matching):
        if index:
            paragraph.add_run("；\n")
        location = evidence.document_name or "来源"
        if evidence.page_number is not None:
            location += f" 第{evidence.page_number}页"
        if evidence.section:
            location += f" · {evidence.section}"
        if evidence.locator:
            location += f" · {evidence.locator}"
        paragraph.add_run(f"{evidence.quote}（{location}） ")
        _add_hyperlink(paragraph, str(evidence.source_url), "原文")


def _notice_kind_label(notice: TenderNotice) -> str:
    return {
        "prequalification": "资格预审公告",
        "tender": "招标公告",
        "correction": "更正公告",
    }.get(notice.opportunity_kind or "", "招标相关公告")


def _budget_label(notice: TenderNotice) -> str:
    if notice.budget is None:
        return "原公告未载明"
    return f"{notice.budget:,.2f} {notice.budget_currency}"


def _notice_tables(document: DocumentType) -> list[Table]:
    required_labels = (
        ("项目名称", "公告类型"),
        ("招标人", "项目编号"),
        ("发布时间", "投标截止"),
        ("项目区域", "预算金额"),
        ("信息来源", "原公告"),
        ("附件", "采集时间"),
    )
    matches = []
    for table in document.tables:
        if len(table.rows) != 6 or len(table.columns) != 4:
            continue
        labels = tuple((row.cells[0].text.strip(), row.cells[2].text.strip()) for row in table.rows)
        if labels == required_labels:
            matches.append(table)
    return matches


def _notice_table_matches(document: DocumentType, table: Table, notice: TenderNotice) -> bool:
    if (
        table.cell(0, 1).text.strip() != notice.title
        or table.cell(2, 1).text.strip() != notice.published_at.strftime("%Y-%m-%d %H:%M")
        or notice.source.source_name not in table.cell(4, 1).text
    ):
        return False
    if _cell_hyperlink_targets(document, table.cell(4, 3)) != {str(notice.source.source_url)}:
        return False
    expected_attachments = {str(item.url) for item in notice.attachments}
    if _cell_hyperlink_targets(document, table.cell(5, 1)) != expected_attachments:
        return False
    if not notice.attachments and table.cell(5, 1).text.strip() != "无":
        return False
    return True


def _cell_hyperlink_targets(document: DocumentType, cell: _Cell) -> set[str]:
    targets: set[str] = set()
    for hyperlink in cell._tc.xpath(".//w:hyperlink"):
        relationship = document.part.rels.get(hyperlink.get(qn("r:id")))
        if relationship is not None and relationship.reltype == RT.HYPERLINK:
            targets.add(relationship.target_ref)
    return targets


def _style_paragraph(
    paragraph: Paragraph,
    size: float,
    color: str,
    *,
    bold: bool = False,
    spacing_before: float = 0,
    spacing_after: float = 0,
) -> None:
    paragraph.paragraph_format.space_before = Pt(spacing_before)
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    for run in paragraph.runs:
        run.font.name = _CJK_FONT
        run.font.size = Pt(size)
        run.font.bold = bold or run.bold
        run.font.color.rgb = RGBColor.from_string(color)
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), _CJK_FONT)
        fonts.set(qn("w:hAnsi"), _CJK_FONT)
        fonts.set(qn("w:eastAsia"), _CJK_FONT)


def _style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = _CJK_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), _CJK_FONT)
    fonts.set(qn("w:hAnsi"), _CJK_FONT)
    fonts.set(qn("w:eastAsia"), _CJK_FONT)


def _add_callout(
    document: DocumentType,
    label: str,
    body: str,
    *,
    fill: str,
    accent: str,
) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_geometry(table, (_TABLE_WIDTH_DXA,))
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)
    _set_cell_margins(cell, 140, 180, 140, 180)
    label_paragraph = _clear_cell(cell)
    label_paragraph.add_run(label)
    _style_paragraph(label_paragraph, 10, accent, bold=True, spacing_after=3)
    body_paragraph = cell.add_paragraph(body or "未生成可展示内容。")
    _style_paragraph(body_paragraph, 10.5, "263238")
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_bullet_group(
    document: DocumentType,
    label: str,
    values: Any,
) -> None:
    items = [str(item).strip() for item in (values or []) if str(item).strip()]
    if not items:
        return
    document.add_heading(label, level=3)
    for item in items:
        paragraph = document.add_paragraph(item, style="List Bullet")
        _style_paragraph(paragraph, 10, "263238", spacing_after=4)


def _set_table_geometry(table: Table, widths_dxa: tuple[int, ...]) -> None:
    if len(widths_dxa) != len(table.columns) or sum(widths_dxa) != _TABLE_WIDTH_DXA:
        raise ValueError("table widths must match columns and sum to 9360 DXA")
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(_TABLE_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(_TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(value / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")


def _style_cell(cell: _Cell, size: float, color: str, *, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        _style_paragraph(paragraph, size, color, bold=bold)


def _shade_cell(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: _Cell, top: int, start: int, bottom: int, end: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _clear_cell(cell: _Cell) -> Paragraph:
    cell.text = ""
    return cell.paragraphs[0]


def _add_hyperlink(paragraph: Paragraph, url: str, label: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((color, underline))
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_number(paragraph: Paragraph) -> None:
    paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, end):
        run = OxmlElement("w:r")
        run.append(element)
        paragraph._p.append(run)
    paragraph.add_run(" 页")
