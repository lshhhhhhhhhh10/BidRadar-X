from __future__ import annotations

from pathlib import Path
from io import BytesIO
import tempfile
import unittest
from unittest.mock import patch
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.api import reports as reports_module
from app.services import publisher as publisher_module
from app.storage import database as database_module
from app.storage.repository import Repository
from app.workflow.nodes import source_select
from tests.integration_support import (
    SuccessfulSource,
    isolated_source_set,
    make_notice,
    source_metadata,
)


class ProductChainApiTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary_directory.name)
        self.report_dir = self.root / "reports"
        self.patches = [
            patch.object(database_module, "DATA_DIR", self.root),
            patch.object(database_module, "DATABASE_PATH", self.root / "app.db"),
            patch.object(publisher_module, "REPORT_DIR", self.report_dir),
            patch.object(reports_module, "REPORT_DIR", self.report_dir),
        ]
        for active_patch in self.patches:
            active_patch.start()
        self.client = TestClient(app)

    def tearDown(self) -> None:
        self.client.close()
        for active_patch in reversed(self.patches):
            active_patch.stop()
        self.temporary_directory.cleanup()

    def test_report_history_is_available_from_public_api(self) -> None:
        response = self.client.get("/api/reports")

        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual(response.json(), {"items": []})

    def test_report_history_delete_persists_and_hides_same_query_duplicates(self) -> None:
        repository = Repository()
        repository.create_task("task-history-delete", "查询全国服务器采购公告", "once")
        base_state = {
            "task_id": "task-history-delete",
            "status": "completed",
            "query": "查询全国服务器采购公告",
            "frequency": "once",
            "projects": [],
            "report": {"status": "no_change"},
        }
        repository.save_run({**base_state, "run_id": "run-history-delete-a"})
        repository.save_run({
            **base_state,
            "run_id": "run-history-delete-b",
            "frequency": "weekly",
        })

        before = self.client.get("/api/reports")
        self.assertEqual(before.status_code, 200, before.text)
        self.assertEqual(len(before.json()["items"]), 1)
        visible_run_id = before.json()["items"][0]["run_id"]

        deleted = self.client.delete(f"/api/reports/{visible_run_id}")
        self.assertEqual(deleted.status_code, 204, deleted.text)

        after_refresh = self.client.get("/api/reports")
        self.assertEqual(after_refresh.status_code, 200, after_refresh.text)
        self.assertEqual(after_refresh.json(), {"items": []})
        with database_module.connect() as connection:
            hidden = connection.execute(
                "SELECT run_id FROM hidden_report_runs ORDER BY run_id"
            ).fetchall()
        self.assertEqual(
            [row["run_id"] for row in hidden],
            ["run-history-delete-a", "run-history-delete-b"],
        )

    def test_completed_run_is_listed_with_a_controlled_docx_download(self) -> None:
        query = f"查询 2026-07-14 全国服务器采购公告 {uuid4()}"
        with patch.object(source_select, "SOURCE_ADAPTERS", isolated_source_set()):
            run_response = self.client.post(
                "/api/tasks/run",
                json={"query": query, "frequency": "once"},
            )

        self.assertEqual(run_response.status_code, 200, run_response.text)
        run = run_response.json()
        history_response = self.client.get("/api/reports")
        self.assertEqual(history_response.status_code, 200, history_response.text)
        item = history_response.json()["items"][0]
        self.assertEqual(item["run_id"], run["run_id"])
        self.assertEqual(item["task_id"], run["task_id"])
        self.assertEqual(item["query"], query)
        self.assertNotEqual(item["display_title"], query)
        self.assertNotIn("-", item["display_title"])
        self.assertEqual(item["report"]["status"], "available")
        self.assertIn(item["ai_report"]["status"], {"generated", "not_generated"})
        self.assertEqual(len(item["projects"]), 1)
        online_project = item["projects"][0]
        self.assertEqual(online_project["title"], "某单位服务器采购公告")
        self.assertEqual(online_project["published_at"], "2026-07-14T09:00:00+08:00")
        self.assertEqual(online_project["url"], "https://public-a.gov.cn/notices/real-001")
        self.assertEqual(online_project["summary"], "采购服务器及配套服务，来源记录 a。")
        self.assertEqual(
            online_project["attachments"][0]["url"],
            "https://public-a.gov.cn/notices/real-001/attachment.pdf",
        )

        download = self.client.get(item["report"]["download_url"])
        self.assertEqual(download.status_code, 200, download.text)
        self.assertEqual(
            download.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment;", download.headers["content-disposition"])
        self.assertIn(".docx", download.headers["content-disposition"])
        document = Document(BytesIO(download.content))
        self.assertTrue(document.paragraphs or document.tables)

    def test_run_projects_and_detail_are_read_from_the_persisted_run(self) -> None:
        with patch.object(source_select, "SOURCE_ADAPTERS", isolated_source_set()):
            run_response = self.client.post(
                "/api/tasks/run",
                json={"query": "查询 2026-07-14 全国服务器采购公告", "frequency": "once"},
            )

        self.assertEqual(run_response.status_code, 200, run_response.text)
        run = run_response.json()
        run_detail = self.client.get(f"/api/runs/{run['run_id']}")
        self.assertEqual(run_detail.status_code, 200, run_detail.text)
        self.assertEqual(run_detail.json()["task_id"], run["task_id"])

        projects_response = self.client.get(f"/api/runs/{run['run_id']}/projects")
        self.assertEqual(projects_response.status_code, 200, projects_response.text)
        projects = projects_response.json()["items"]
        self.assertEqual(len(projects), 1)
        project = projects[0]
        self.assertEqual(project["title"], "某单位服务器采购公告")
        self.assertFalse(project["details_loaded"])
        self.assertEqual(project["module_count"], 0)
        self.assertEqual(
            project["url"],
            "https://public-a.gov.cn/notices/real-001",
        )
        self.assertEqual(project["attachments"][0]["name"], "采购需求附件.pdf")

        detail_response = self.client.get(
            f"/api/runs/{run['run_id']}/projects/{project['project_id']}"
        )
        self.assertEqual(detail_response.status_code, 200, detail_response.text)
        detail = detail_response.json()
        self.assertEqual(detail["run_id"], run["run_id"])
        self.assertTrue(detail["details_loaded"])
        self.assertEqual(detail["summary"], "采购服务器及配套服务，来源记录 a。")
        self.assertEqual(detail["modules"], [])
        self.assertEqual(
            self.client.get(
                f"/api/runs/{run['run_id']}/projects/project-does-not-exist"
            ).status_code,
            404,
        )

    def test_run_report_endpoint_exposes_the_current_available_report(self) -> None:
        with patch.object(source_select, "SOURCE_ADAPTERS", isolated_source_set()):
            run_response = self.client.post(
                "/api/tasks/run",
                json={"query": f"查询服务器采购 {uuid4()}", "frequency": "once"},
            )

        run = run_response.json()
        report_response = self.client.get(f"/api/runs/{run['run_id']}/report")
        self.assertEqual(report_response.status_code, 200, report_response.text)
        report = report_response.json()
        self.assertEqual(report["status"], "available")
        self.assertEqual(
            report["download_url"],
            f"/api/reports/{report['delivery_fingerprint']}/download",
        )

    def test_each_changed_project_has_an_independent_downloadable_word_file(self) -> None:
        sources = isolated_source_set()
        second_notice = make_notice(
            source_id="public-a",
            source_name="公开来源 A",
            source_url="https://public-a.gov.cn/notices/storage-002",
            marker="c",
            project_fingerprint="9" * 64,
        ).model_copy(
            update={
                "notice_id": "public-a-storage-002",
                "title": "新增存储设备采购公告",
                "notice_stable_fingerprint": "8" * 64,
                "raw_content_fingerprint": "7" * 64,
            }
        )
        sources[0].notices.append(second_notice)
        with patch.object(source_select, "SOURCE_ADAPTERS", sources):
            response = self.client.post(
                "/api/tasks/run",
                json={"query": f"查询服务器和存储设备采购 {uuid4()}", "frequency": "daily"},
            )

        self.assertEqual(response.status_code, 200, response.text)
        run = response.json()
        report = self.client.get(f"/api/runs/{run['run_id']}/report").json()
        self.assertEqual(report["status"], "available")
        self.assertEqual(report["document_count"], 2)
        self.assertEqual(len(report["documents"]), 2)
        self.assertEqual(len({item["download_url"] for item in report["documents"]}), 2)
        self.assertTrue(all(item["is_new"] for item in report["documents"]))
        self.assertTrue(all(item["status"] == "available" for item in report["documents"]))

        by_title = {item["project_title"]: item for item in report["documents"]}
        for title, document_view in by_title.items():
            download = self.client.get(document_view["download_url"])
            self.assertEqual(download.status_code, 200, download.text)
            document = Document(BytesIO(download.content))
            text = "\n".join(
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            self.assertIn(title, text)
            other_titles = set(by_title) - {title}
            self.assertTrue(all(other_title not in text for other_title in other_titles))

    def test_failed_task_returns_a_structured_error_without_internal_details(self) -> None:
        with (
            patch.object(source_select, "SOURCE_ADAPTERS", isolated_source_set()),
            patch.object(
                publisher_module.DocxPublisher,
                "publish",
                side_effect=RuntimeError(r"failed at C:\secret\app.db API_KEY=value"),
            ),
        ):
            response = self.client.post(
                "/api/tasks/run",
                json={"query": f"查询服务器失败任务 {uuid4()}", "frequency": "once"},
            )

        self.assertEqual(response.status_code, 502, response.text)
        self.assertEqual(response.json()["detail"]["code"], "task_failed")
        public_text = response.text.lower()
        self.assertNotIn("app.db", public_text)
        self.assertNotIn("api_key", public_text)
        self.assertNotIn("traceback", public_text)

        history = self.client.get("/api/reports").json()["items"]
        self.assertEqual(history[0]["run_status"], "failed")
        self.assertEqual(history[0]["report"]["status"], "failed")
        self.assertNotIn("app.db", str(history[0]).lower())

        run_history = self.client.get("/api/runs")
        self.assertEqual(run_history.status_code, 200, run_history.text)
        self.assertNotIn("app.db", run_history.text.lower())
        self.assertNotIn("api_key", run_history.text.lower())

    def test_partial_source_failure_does_not_expose_adapter_error_details(self) -> None:
        sources = isolated_source_set()
        sources[-1].message = r"failed at C:\secret\source.db TENDER_DATA_DIR=C:\secret"
        with patch.object(source_select, "SOURCE_ADAPTERS", sources):
            response = self.client.post(
                "/api/tasks/run",
                json={"query": f"查询部分来源失败 {uuid4()}", "frequency": "once"},
            )

        self.assertEqual(response.status_code, 200, response.text)
        public_text = response.text.lower()
        self.assertNotIn("source.db", public_text)
        self.assertNotIn("tender_data_dir", public_text)
        failed_sources = [
            source
            for source in response.json()["selected_sources"]
            if source["collection_status"] == "failed"
        ]
        self.assertEqual(len(failed_sources), 1)
        self.assertEqual(failed_sources[0]["error"], "来源采集失败，请稍后重试。")

    def test_empty_run_does_not_invent_projects_or_urls(self) -> None:
        empty_source = SuccessfulSource(source_metadata("empty", "空结果公开来源"), [])
        with patch.object(source_select, "SOURCE_ADAPTERS", [empty_source]):
            run_response = self.client.post(
                "/api/tasks/run",
                json={"query": "查询不存在的真实公告条件", "frequency": "once"},
            )

        self.assertEqual(run_response.status_code, 200, run_response.text)
        run = run_response.json()
        self.assertEqual(run["projects"], [])
        projects = self.client.get(f"/api/runs/{run['run_id']}/projects")
        self.assertEqual(projects.status_code, 200, projects.text)
        self.assertEqual(projects.json(), {"items": []})

    def test_missing_run_and_project_are_not_reported_as_empty_successes(self) -> None:
        self.assertEqual(self.client.get("/api/runs/missing").status_code, 404)
        self.assertEqual(self.client.get("/api/runs/missing/projects").status_code, 404)
        self.assertEqual(self.client.get("/api/runs/missing/report").status_code, 404)
        self.assertEqual(
            self.client.get("/api/runs/missing/projects/project-missing").status_code,
            404,
        )

    def test_missing_docx_and_invalid_download_identifier_have_distinct_errors(self) -> None:
        with patch.object(source_select, "SOURCE_ADAPTERS", isolated_source_set()):
            run = self.client.post(
                "/api/tasks/run",
                json={"query": f"查询待删除服务器报告 {uuid4()}", "frequency": "once"},
            ).json()
        report = self.client.get(f"/api/runs/{run['run_id']}/report").json()
        download_url = report["download_url"]
        next(self.report_dir.glob("*.docx")).unlink()

        missing = self.client.get(f"/api/runs/{run['run_id']}/report")
        self.assertEqual(missing.status_code, 200, missing.text)
        self.assertEqual(missing.json()["status"], "missing")
        self.assertEqual(self.client.get(download_url).status_code, 410)
        invalid = self.client.get("/api/reports/not-a-fingerprint/download")
        self.assertEqual(invalid.status_code, 400, invalid.text)
        self.assertNotIn("app.db", invalid.text.lower())
        traversal = self.client.get("/api/reports/..%5Capp.db/download")
        self.assertIn(traversal.status_code, {400, 404})
        self.assertNotIn("sqlite", traversal.text.lower())
        self.assertEqual(
            self.client.get(f"/api/reports/{'f' * 64}/download").status_code,
            404,
        )

    def test_no_change_run_is_distinct_from_failed_and_missing_reports(self) -> None:
        query = f"查询重复运行报告 {uuid4()}"
        with patch.object(source_select, "SOURCE_ADAPTERS", isolated_source_set()):
            first = self.client.post(
                "/api/tasks/run", json={"query": query, "frequency": "once"}
            )
            second = self.client.post(
                "/api/tasks/run", json={"query": query, "frequency": "once"}
            )

        self.assertEqual(first.status_code, 200, first.text)
        self.assertEqual(second.status_code, 200, second.text)
        report = self.client.get(
            f"/api/runs/{second.json()['run_id']}/report"
        )
        self.assertEqual(report.status_code, 200, report.text)
        self.assertEqual(report.json()["status"], "not_generated")


if __name__ == "__main__":
    unittest.main()
