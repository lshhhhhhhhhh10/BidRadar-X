from datetime import datetime
from pathlib import Path
import re
import tempfile
import unittest
from urllib.parse import unquote

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from fastapi.testclient import TestClient

from app.api import reports as reports_api
from app.main import app
from app.services.demo_shanghai_property import DEMO_QUERY, demo_notices
from app.services.docx_publisher import DocxPublisher


def document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    return "\n".join(parts)


class ShanghaiPropertyDemoReportTest(unittest.TestCase):
    def test_scheduled_run_download_contains_only_that_runs_new_notices(self) -> None:
        with tempfile.TemporaryDirectory() as output_dir:
            original_report_dir = reports_api.REPORT_DIR
            reports_api.REPORT_DIR = Path(output_dir)
            try:
                response = TestClient(app).get(
                    "/api/demo/reports/shanghai-property/runs/20260715-0900/download"
                )
            finally:
                reports_api.REPORT_DIR = original_report_dir

        self.assertEqual(response.status_code, 200)
        self.assertIn(
            "上海市物业管理服务项目_202607150900.docx",
            unquote(response.headers["content-disposition"]),
        )
        with tempfile.TemporaryDirectory() as readback_dir:
            report_path = Path(readback_dir) / "incremental.docx"
            report_path.write_bytes(response.content)
            text = document_text(Document(report_path))
        self.assertIn("老港项目设备维修、维护服务招标公告", text)
        self.assertIn("北艾路1400号物业服务费的公开招标公告", text)
        self.assertNotIn("物业客服服务和物业设施日常维护服务公开招标公告", text)

    def test_exact_demo_query_generates_ten_project_report_with_eight_modules(self) -> None:
        notices = demo_notices()
        self.assertEqual(len(notices), 10)
        self.assertTrue(
            all(str(notice.source.source_url).startswith("https://example.invalid/") for notice in notices)
        )
        self.assertTrue(all(notice.source.authority == 0 for notice in notices))
        self.assertTrue(all(notice.title.startswith("[合成演示]") for notice in notices))
        self.assertEqual(
            [section.section_id for section in notices[0].requirement_sections],
            [
                "procurement", "qualification", "technical", "timeline",
                "commercial", "submission", "evaluation", "reference",
            ],
        )
        known_fact_count = sum(
            len([fact for fact in section.facts if fact.value is not None])
            for section in notices[0].requirement_sections
        )
        self.assertEqual(known_fact_count, 72)

        generated_at = datetime.fromisoformat("2026-07-15T17:26:00+08:00")
        with tempfile.TemporaryDirectory() as output_dir:
            report_path = DocxPublisher(
                output_dir=Path(output_dir),
                clock=lambda: generated_at,
            ).publish(DEMO_QUERY, notices, report_scope="full")

            self.assertRegex(
                report_path.name,
                re.compile(r"^上海市物业管理服务项目_\d{12}\.docx$"),
            )
            document = Document(report_path)
            text = document_text(document)
            for notice in notices:
                self.assertIn(notice.title, text)
            for heading in (
                "项目及采购内容", "投标人资格要求", "技术与服务要求",
                "项目周期与验收要求", "报价、付款与保证金",
                "投标组织与文件要求", "评标与定标规则", "客观参考信息",
            ):
                self.assertIn(heading, text)
            self.assertIn("采购数量或项目规模", text)
            self.assertIn("服务建筑面积约 12.8 万平方米。", text)

            targets = {
                relationship.target_ref
                for relationship in document.part.rels.values()
                if relationship.reltype == RT.HYPERLINK
            }
            self.assertTrue(
                {str(notice.source.source_url) for notice in notices}.issubset(targets)
            )


if __name__ == "__main__":
    unittest.main()
