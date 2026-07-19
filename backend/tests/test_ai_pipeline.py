from __future__ import annotations

from io import BytesIO
import json
import os
from pathlib import Path
import tempfile
import unittest
from unittest.mock import patch
from urllib.error import HTTPError

from docx import Document
from fastapi.testclient import TestClient

from app.ai.client import StructuredAIClient
from app.ai.config import (
    AISettings,
    clear_runtime_api_key,
    clear_runtime_profiles,
    configure_runtime_profile,
)
from app.ai.prompts import INTENT_PROMPT
from app.main import app
from app.services.docx_publisher import DocxPublisher
from tests.integration_support import make_notice


class AIPipelineTest(unittest.TestCase):
    def tearDown(self) -> None:
        clear_runtime_api_key()
        clear_runtime_profiles()

    def test_structured_client_validates_json_and_audits_without_secret(self) -> None:
        captured: dict[str, object] = {}

        def transport(request, timeout):
            captured["timeout"] = timeout
            captured["authorization"] = request.get_header("Authorization")
            captured["url"] = request.full_url
            payload = json.loads(request.data.decode("utf-8"))
            captured["response_format"] = payload["response_format"]
            captured["system_message"] = payload["messages"][0]["content"]
            return {
                "id": "resp-test",
                "choices": [
                    {
                        "message": {
                            "role": "assistant",
                            "content": json.dumps(
                                    {
                                        "topic": "服务器",
                                        "regions": ["上海市"],
                                        "keywords": ["服务器", "计算节点"],
                                        "exclusions": [],
                                        "time_range_start": None,
                                        "time_range_end": None,
                                        "confidence": 0.96,
                                        "interpretation": "检索上海服务器采购公告",
                                    },
                                    ensure_ascii=False,
                                ),
                        }
                    }
                ],
                "usage": {"prompt_tokens": 100, "completion_tokens": 50},
            }

        secret = "test-secret-this-must-never-enter-audit"
        settings = AISettings(
            provider="zhipu",
            api_key=secret,
            base_url="https://open.bigmodel.cn/api/paas/v4",
            model="glm-test",
            protocol="chat_completions_json",
            timeout_seconds=12,
            enabled=True,
        )
        result = StructuredAIClient(settings=settings, transport=transport).complete(
            INTENT_PROMPT,
            {"query": "上海服务器采购"},
        )

        self.assertIsNotNone(result.value)
        self.assertEqual(result.value.topic, "服务器")
        self.assertEqual(result.audit["status"], "completed")
        self.assertEqual(result.audit["input_tokens"], 100)
        self.assertNotIn(secret, json.dumps(result.audit, ensure_ascii=False))
        self.assertEqual(captured["authorization"], f"Bearer {secret}")
        self.assertEqual(captured["response_format"], {"type": "json_object"})
        self.assertIn("time_range_start", captured["system_message"])
        self.assertEqual(
            captured["url"],
            "https://open.bigmodel.cn/api/paas/v4/chat/completions",
        )

    def test_invalid_model_output_returns_safe_fallback_signal(self) -> None:
        settings = AISettings(
            provider="zhipu",
            api_key="sk-test-invalid-output",
            base_url="https://open.bigmodel.cn/api/paas/v4",
            model="glm-test",
            protocol="chat_completions_json",
            timeout_seconds=12,
            enabled=True,
        )
        client = StructuredAIClient(
            settings=settings,
            transport=lambda request, timeout: {
                "id": "resp-invalid",
                "choices": [{"message": {"content": "not-json"}}],
            },
        )

        result = client.complete(INTENT_PROMPT, {"query": "服务器"})

        self.assertIsNone(result.value)
        self.assertEqual(result.audit["status"], "invalid_output")
        self.assertNotIn("not-json", json.dumps(result.audit))

    def test_balance_error_falls_back_to_free_model_and_keeps_safe_audit(self) -> None:
        requested_models: list[str] = []

        def transport(request, timeout):
            payload = json.loads(request.data.decode("utf-8"))
            requested_models.append(payload["model"])
            if payload["model"] == "glm-paid":
                raise HTTPError(
                    request.full_url,
                    429,
                    "balance",
                    {},
                    BytesIO(json.dumps({"error": {"code": "1113"}}).encode()),
                )
            return {
                "id": "resp-free",
                "choices": [{"message": {"content": json.dumps({
                    "topic": "服务器",
                    "regions": ["全国"],
                    "keywords": ["服务器"],
                    "exclusions": [],
                    "time_range_start": None,
                    "time_range_end": None,
                    "confidence": 0.9,
                    "interpretation": "检索全国服务器采购公告",
                }, ensure_ascii=False)}}],
                "usage": {"prompt_tokens": 10, "completion_tokens": 8},
            }

        settings = AISettings(
            provider="zhipu",
            api_key="test-secret",
            base_url="https://open.bigmodel.cn/api/paas/v4",
            model="glm-paid",
            protocol="chat_completions_json",
            timeout_seconds=12,
            enabled=True,
            fallback_model="glm-4.7-flash",
        )
        result = StructuredAIClient(settings=settings, transport=transport).complete(
            INTENT_PROMPT,
            {"query": "全国服务器采购"},
        )

        self.assertIsNotNone(result.value)
        self.assertEqual(requested_models, ["glm-paid", "glm-4.7-flash"])
        self.assertEqual(result.audit["status"], "completed")
        self.assertEqual(result.audit["model"], "glm-4.7-flash")
        self.assertEqual(result.audit["fallback_from"], "glm-paid")

    def test_multiple_backend_profiles_fail_over_without_exposing_either_key(self) -> None:
        first_key = "first-provider-secret"
        second_key = "second-provider-secret"
        configure_runtime_profile(
            label="主密钥",
            provider="zhipu",
            api_key=first_key,
            model="glm-primary",
            priority=0,
        )
        configure_runtime_profile(
            label="备用密钥",
            provider="deepseek",
            api_key=second_key,
            model="deepseek-chat",
            priority=10,
        )
        authorizations: list[str] = []

        def transport(request, timeout):
            del timeout
            authorization = request.get_header("Authorization")
            authorizations.append(authorization)
            if authorization == f"Bearer {first_key}":
                raise HTTPError(
                    request.full_url,
                    429,
                    "balance",
                    {},
                    BytesIO(json.dumps({"error": {"code": "1113"}}).encode()),
                )
            return {
                "id": "resp-failover",
                "choices": [{"message": {"content": json.dumps({
                    "topic": "服务器",
                    "regions": ["全国"],
                    "keywords": ["服务器"],
                    "exclusions": [],
                    "time_range_start": None,
                    "time_range_end": None,
                    "confidence": 0.9,
                    "interpretation": "检索全国服务器采购公告",
                }, ensure_ascii=False)}}],
            }

        with patch.dict(os.environ, {"BIDRADAR_AI_API_KEY": ""}, clear=False):
            result = StructuredAIClient(transport=transport).complete(
                INTENT_PROMPT,
                {"query": "服务器"},
            )

        self.assertIsNotNone(result.value)
        self.assertEqual(authorizations, [f"Bearer {first_key}", f"Bearer {second_key}"])
        self.assertEqual(result.audit["provider"], "deepseek")
        self.assertEqual(result.audit["failover_count"], 1)
        serialized = json.dumps(result.audit, ensure_ascii=False)
        self.assertNotIn(first_key, serialized)
        self.assertNotIn(second_key, serialized)

    def test_rate_limited_profile_immediately_fails_over_and_enters_backoff(self) -> None:
        first_key = "rate-limited-profile-secret"
        second_key = "healthy-profile-secret"
        configure_runtime_profile(
            label="限流密钥",
            provider="zhipu",
            api_key=first_key,
            model="glm-primary",
            priority=0,
        )
        configure_runtime_profile(
            label="健康密钥",
            provider="zhipu",
            api_key=second_key,
            model="glm-secondary",
            priority=10,
        )
        authorizations: list[str] = []

        def transport(request, timeout):
            del timeout
            authorization = request.get_header("Authorization")
            authorizations.append(authorization)
            if authorization == f"Bearer {first_key}":
                raise HTTPError(
                    request.full_url,
                    429,
                    "rate-limit",
                    {},
                    BytesIO(json.dumps({"error": {"code": "1302"}}).encode()),
                )
            return {
                "id": "resp-after-rate-limit",
                "choices": [{"message": {"content": json.dumps({
                    "topic": "服务器",
                    "regions": ["全国"],
                    "keywords": ["服务器"],
                    "exclusions": [],
                    "time_range_start": None,
                    "time_range_end": None,
                    "confidence": 0.9,
                    "interpretation": "检索全国服务器采购公告",
                }, ensure_ascii=False)}}],
            }

        with patch.dict(
            os.environ,
            {
                "BIDRADAR_AI_API_KEY": "",
                "BIDRADAR_AI_SECONDARY_API_KEY": "",
            },
            clear=False,
        ):
            client = StructuredAIClient(transport=transport)
            first_result = client.complete(INTENT_PROMPT, {"query": "服务器"})
            second_result = client.complete(INTENT_PROMPT, {"query": "服务器"})

        self.assertIsNotNone(first_result.value)
        self.assertIsNotNone(second_result.value)
        self.assertEqual(
            authorizations,
            [
                f"Bearer {first_key}",
                f"Bearer {second_key}",
                f"Bearer {second_key}",
            ],
        )
        self.assertEqual(first_result.audit["failover_count"], 1)
        self.assertEqual(second_result.audit["status"], "completed")

    def test_environment_secondary_key_is_a_persistent_failover_candidate(self) -> None:
        primary_key = "primary-environment-secret"
        secondary_key = "secondary-environment-secret"
        with patch.dict(
            os.environ,
            {
                "BIDRADAR_AI_API_KEY": primary_key,
                "BIDRADAR_AI_SECONDARY_API_KEY": secondary_key,
                "BIDRADAR_AI_PROVIDER": "zhipu",
                "BIDRADAR_AI_MODEL": "glm-primary",
                "BIDRADAR_AI_SECONDARY_MODEL": "glm-5.2",
                "BIDRADAR_AI_ENABLED": "auto",
            },
            clear=False,
        ):
            candidates = AISettings.candidates()

        environment_candidates = [
            candidate
            for candidate in candidates
            if candidate.profile_id.startswith("environment")
        ]
        self.assertEqual(
            [(item.profile_id, item.model) for item in environment_candidates],
            [("environment", "glm-primary"), ("environment-secondary", "glm-5.2")],
        )
        serialized = json.dumps(
            [item.public_status() for item in environment_candidates],
            ensure_ascii=False,
        )
        self.assertNotIn(primary_key, serialized)
        self.assertNotIn(secondary_key, serialized)

    def test_profile_api_manages_multiple_redacted_backend_credentials(self) -> None:
        with patch.dict(os.environ, {"BIDRADAR_AI_API_KEY": ""}, clear=False), TestClient(app) as client:
            secret = "profile-api-secret-value"
            created = client.post("/api/ai/profiles", json={
                "label": "DeepSeek 备用",
                "provider": "deepseek",
                "api_key": secret,
                "model": "deepseek-chat",
            })
            self.assertEqual(created.status_code, 201, created.text)
            self.assertNotIn(secret, created.text)
            profile_id = created.json()["profile"]["profile_id"]
            listed = client.get("/api/ai/profiles")
            self.assertEqual(listed.status_code, 200)
            self.assertEqual(len(listed.json()["items"]), 1)
            self.assertNotIn(secret, listed.text)
            disabled = client.patch(
                f"/api/ai/profiles/{profile_id}",
                json={"enabled": False, "priority": 30},
            )
            self.assertFalse(disabled.json()["profile"]["enabled"])
            self.assertEqual(client.delete(f"/api/ai/profiles/{profile_id}").status_code, 204)

    def test_local_backend_credential_enables_status_without_returning_key(self) -> None:
        with patch.dict(os.environ, {"BIDRADAR_AI_API_KEY": ""}, clear=False), TestClient(app) as client:
            before = client.get("/api/ai/status")
            self.assertEqual(before.status_code, 200)
            self.assertFalse(before.json()["enabled"])

            secret = "test-secret-local-backend-memory-only"
            connected = client.put("/api/ai/credential", json={"api_key": secret})
            self.assertEqual(connected.status_code, 200, connected.text)
            self.assertTrue(connected.json()["enabled"])
            self.assertNotIn(secret, connected.text)

            disconnected = client.delete("/api/ai/credential")
            self.assertEqual(disconnected.status_code, 200)
            self.assertFalse(disconnected.json()["enabled"])

    def test_docx_keeps_original_content_and_adds_evidenced_ai_section(self) -> None:
        notice = make_notice(
            source_id="public-a",
            source_name="公开来源 A",
            source_url="https://public-a.gov.cn/notices/real-001",
            marker="a",
        )
        ai_report = {
            "status": "generated",
            "executive_summary": "发现一条与服务器采购直接相关的公告。",
            "key_findings": [
                {"text": "公告采购服务器及配套服务。", "evidence_ids": ["ev-test"]}
            ],
            "notice_narratives": [
                {
                    "notice_id": notice.notice_id,
                    "summary": "建议核对采购需求附件与截止时间。",
                    "risk_level": "medium",
                    "risk_assessment": "截止时间明确，但资格与附件要求仍需逐项复核。",
                    "risk_points": ["截止时间需以原公告为准"],
                    "opportunity_points": ["采购内容与服务器供应能力直接相关"],
                    "next_actions": ["打开原公告复核"],
                    "evidence_ids": ["ev-test"],
                }
            ],
        }
        with tempfile.TemporaryDirectory() as output_dir:
            report_path = DocxPublisher(Path(output_dir)).publish(
                "服务器采购",
                [notice],
                ai_report=ai_report,
            )
            document = Document(BytesIO(report_path.read_bytes()))

        all_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                paragraph.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            ]
        )
        self.assertIn("AI 辅助摘要与风险研判", all_text)
        self.assertIn("建议核对采购需求附件与截止时间", all_text)
        self.assertIn("风险等级", all_text)
        self.assertIn("机会提示", all_text)
        self.assertIn(notice.core_content, all_text)


if __name__ == "__main__":
    unittest.main()
