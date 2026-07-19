from datetime import datetime
import json
from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from app.schemas.tender import TenderNotice
from app.services.docx_publisher import (
    DocxPublisher,
    DocxValidationError,
    build_report_filename,
    validate_docx,
)


FIXTURE_PATH = Path(__file__).parent / "fixtures" / "reports" / "tender_notices.json"


def load_notices() -> list[TenderNotice]:
    payload = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))
    return [TenderNotice.model_validate(item) for item in payload]


def all_document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    return "\n".join(parts)


class DocxPublisherTest(unittest.TestCase):
    def test_public_filename_uses_only_question_time_and_optional_project_sequence(self) -> None:
        generated_at = datetime.fromisoformat("2026-07-14T12:30:45+08:00")
        self.assertEqual(
            build_report_filename("最近3个月的服务器采购", generated_at),
            "最近3个月的服务器采购_202607141230.docx",
        )
        self.assertEqual(
            build_report_filename("最近3个月的服务器采购", generated_at, project_sequence=2),
            "最近3个月的服务器采购_202607141230_项目02.docx",
        )
    def test_full_report_preserves_chinese_and_renders_required_content_and_links(self) -> None:
        notices = load_notices()
        generated_at = datetime.fromisoformat("2026-07-14T12:30:45+08:00")

        with tempfile.TemporaryDirectory() as output_dir:
            report_path = DocxPublisher(
                output_dir=Path(output_dir),
                clock=lambda: generated_at,
            ).publish(
                query='安徽省/服务器:*采购?<>"',
                notices=notices,
                report_scope="full",
            )

            self.assertEqual(report_path.name, "安徽省_服务器_采购_202607141230.docx")
            document = Document(report_path)
            text = all_document_text(document)
            self.assertIn("全量报告", text)
            for notice in notices:
                self.assertIn(notice.title, text)
                self.assertIn(notice.core_content, text)
            self.assertIn("无", text)
            self.assertIn("技术参数矩阵", text)
            self.assertIn("技术要求", text)
            self.assertIn("设备技术参数以采购需求文件为准。", text)
            self.assertGreaterEqual(len(document.tables), 4)

            targets = {
                relationship.target_ref
                for relationship in document.part.rels.values()
                if relationship.reltype == RT.HYPERLINK
            }
            expected_targets = {
                str(notice.source.source_url) for notice in notices
            } | {
                str(attachment.url)
                for notice in notices
                for attachment in notice.attachments
            }
            self.assertTrue(expected_targets.issubset(targets))

    def test_incremental_report_contains_only_supplied_new_notices_and_never_overwrites(self) -> None:
        notices = load_notices()
        generated_at = datetime.fromisoformat("2026-07-14T12:31:59+08:00")

        with tempfile.TemporaryDirectory() as output_dir:
            publisher = DocxPublisher(
                output_dir=Path(output_dir),
                clock=lambda: generated_at,
            )
            report_path = publisher.publish(
                query="安徽省服务器采购",
                notices=notices[:1],
                report_scope="incremental",
            )
            original_bytes = report_path.read_bytes()

            document = Document(report_path)
            text = all_document_text(document)
            self.assertIn("仅新增内容报告", text)
            self.assertIn(notices[0].title, text)
            self.assertNotIn(notices[1].title, text)

            with self.assertRaisesRegex(FileExistsError, "already exists"):
                publisher.publish(
                    query="安徽省服务器采购",
                    notices=notices[:1],
                    report_scope="incremental",
                )
            self.assertEqual(report_path.read_bytes(), original_bytes)

    def test_empty_report_is_valid_and_has_no_invented_notice_or_attachment(self) -> None:
        generated_at = datetime.fromisoformat("2026-07-14T12:32:00+08:00")

        with tempfile.TemporaryDirectory() as output_dir:
            report_path = DocxPublisher(
                output_dir=Path(output_dir),
                clock=lambda: generated_at,
            ).publish(query=".", notices=[], report_scope="full")

            self.assertEqual(report_path.name, "情报任务_202607141232.docx")
            document = Document(report_path)
            text = all_document_text(document)
            self.assertIn("未发现符合条件的公告。", text)
            self.assertNotIn("附件链接", text)

            validation = validate_docx(
                report_path,
                expected_notices=[],
                expected_scope="full",
            )
            self.assertGreater(validation.paragraph_count, 0)
            self.assertEqual(validation.table_count, 2)
            self.assertEqual(validation.hyperlink_targets, ())

    def test_validation_requires_a_distinct_fact_table_for_each_notice(self) -> None:
        notice = load_notices()[1]
        same_content_notice = notice.model_copy(
            update={
                "notice_id": "fixture-notice-002-copy",
                "raw_content_fingerprint": "1" * 64,
                "notice_stable_fingerprint": "2" * 64,
                "project_stable_fingerprint": "3" * 64,
            }
        )
        generated_at = datetime.fromisoformat("2026-07-14T12:33:00+08:00")

        with tempfile.TemporaryDirectory() as output_dir:
            report_path = DocxPublisher(
                output_dir=Path(output_dir),
                clock=lambda: generated_at,
            ).publish(query="重复标题验证", notices=[notice], report_scope="full")

            with self.assertRaisesRegex(DocxValidationError, "notice table"):
                validate_docx(
                    report_path,
                    expected_notices=[notice, same_content_notice],
                    expected_scope="full",
                )


if __name__ == "__main__":
    unittest.main()
