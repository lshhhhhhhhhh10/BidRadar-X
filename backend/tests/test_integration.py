import asyncio
import json
from pathlib import Path
import tempfile
import unittest
from unittest.mock import patch
from uuid import uuid4

from docx import Document

from app.schemas.tender import TaskSpec
from app.services import publisher as publisher_module
from app.sources import build_production_sources
from app.workflow.graph import WORKFLOW
from app.workflow.nodes import source_select
from tests.integration_support import FailingSource, isolated_source_set, source_metadata


class RealWorkflowIntegrationTest(unittest.TestCase):
    def test_production_registry_contains_real_public_and_login_sources_only(self) -> None:
        with patch.dict(
            "os.environ",
            {"BIDRADAR_TIANYANCHA_TOKEN": "configured-for-registry-test"},
        ):
            sources = build_production_sources()

        self.assertGreaterEqual(len(sources), 2)
        self.assertTrue(any(item.metadata.get("requires_login") for item in sources))
        self.assertTrue(any(not item.metadata.get("requires_login") for item in sources))
        self.assertNotIn("example.local", repr(sources))
        self.assertNotIn("public-platform", {item.metadata["source_id"] for item in sources})
        self.assertIn("tianyancha-bids", {item.metadata["source_id"] for item in sources})

    def test_failed_login_source_does_not_block_real_docx_report(self) -> None:
        adapters = isolated_source_set()
        with tempfile.TemporaryDirectory() as output_dir:
            with (
                patch.object(source_select, "SOURCE_ADAPTERS", adapters),
                patch.object(publisher_module, "REPORT_DIR", Path(output_dir)),
            ):
                state = asyncio.run(
                    WORKFLOW.ainvoke(
                        {
                            "task_id": str(uuid4()),
                            "run_id": str(uuid4()),
                            "query": "查询 2026-07-14 全国服务器采购公告",
                            "frequency": "once",
                            "status": "running",
                            "steps": [],
                            "funnel": {},
                            "retry_count": 0,
                            "quality_passed": False,
                            "quality_issues": [],
                        },
                        config={"recursion_limit": 50},
                    )
                )

            TaskSpec.model_validate(state["task_spec"])
            report = state["report"]
            self.assertEqual(state["status"], "completed")
            self.assertEqual(report["source_count"], 3)
            self.assertEqual(len(report["successful_sources"]), 2)
            self.assertEqual(len(report["failed_sources"]), 1)
            self.assertEqual(report["failed_sources"][0]["source_id"], "login-source")
            self.assertTrue(report["filename"].endswith(".docx"))
            self.assertEqual(len(state["projects"]), 1)

            report_path = Path(output_dir) / report["filename"]
            self.assertTrue(report_path.is_file())
            document_text = "\n".join(
                paragraph.text
                for table in Document(report_path).tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            self.assertIn("某单位服务器采购公告", document_text)
            self.assertIn("https://public-a.gov.cn/notices/real-001", document_text)
            self.assertIn("https://public-b.gov.cn/notices/real-001", document_text)
            self.assertNotIn("example.local", document_text)
            self.assertNotIn("模拟证据", document_text)
            self.assertNotIn("example.local", json.dumps(state, ensure_ascii=False))

    def test_all_failed_sources_mark_the_run_failed_without_creating_a_report(self) -> None:
        adapters = [
            FailingSource(source_metadata("public-a", "公开来源 A")),
            FailingSource(
                source_metadata("login-source", "登录来源", requires_login=True)
            ),
        ]
        with tempfile.TemporaryDirectory() as output_dir:
            with (
                patch.object(source_select, "SOURCE_ADAPTERS", adapters),
                patch.object(publisher_module, "REPORT_DIR", Path(output_dir)),
            ):
                state = asyncio.run(
                    WORKFLOW.ainvoke(
                        {
                            "task_id": str(uuid4()),
                            "run_id": str(uuid4()),
                            "query": "查询 2026-07-14 全国服务器采购公告",
                            "frequency": "once",
                            "status": "running",
                            "steps": [],
                            "funnel": {},
                            "retry_count": 0,
                            "quality_passed": False,
                            "quality_issues": [],
                        },
                        config={"recursion_limit": 50},
                    )
                )

            self.assertEqual(state["status"], "failed")
            self.assertNotIn("report", state)
            self.assertEqual(
                len([item for item in state["selected_sources"] if item["collection_status"] == "failed"]),
                2,
            )
            self.assertEqual(list(Path(output_dir).glob("*.docx")), [])


if __name__ == "__main__":
    unittest.main()
