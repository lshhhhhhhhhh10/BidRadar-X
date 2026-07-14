from __future__ import annotations

import asyncio
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
import os
from pathlib import Path
import tempfile
import time
import unittest
from unittest.mock import patch
from uuid import uuid4

from docx import Document

from app.schemas.tender import EvidenceReference, SourceRecord, TenderNotice
from app.services import publisher as publisher_module
from app.services.publisher import Publisher
from app.intelligence.change_detector import ChangeDetector
from app.storage import database as database_module
from app.storage.repository import Repository
from app.workflow.graph import WORKFLOW
from app.workflow.nodes import source_select
from tests.integration_support import FailingSource, SuccessfulSource, source_metadata


def make_notice(
    *,
    source_id: str = "public-a",
    source_url: str = "https://public-a.gov.cn/notices/001",
    project_fingerprint: str = "1" * 64,
    notice_fingerprint: str = "2" * 64,
    raw_fingerprint: str = "3" * 64,
    notice_type: str = "tender",
    budget: str = "1000000.00",
    deadline: str = "2026-07-31T17:00:00+08:00",
    purchaser: str = "某采购单位",
    fetched_at: str = "2026-07-14T10:00:00+08:00",
    title: str = "某单位服务器采购公告",
    published_at: str = "2026-07-14T09:00:00+08:00",
) -> TenderNotice:
    evidence = [
        EvidenceReference(
            evidence_id=f"evidence-{field_path}",
            field_path=field_path,
            source_url=source_url,
            quote=f"{field_path} 的公告原文证据",
            fetched_at=fetched_at,
        )
        for field_path in ("budget", "deadline", "purchaser")
    ]
    return TenderNotice(
        notice_id=f"{source_id}-{raw_fingerprint[:8]}",
        notice_type=notice_type,
        title=title,
        published_at=published_at,
        source=SourceRecord(
            source_id=source_id,
            source_name=f"公开来源 {source_id}",
            source_url=source_url,
            publication_role="original",
            source_notice_id=notice_fingerprint[:12],
            authority=1.0,
        ),
        core_content="采购服务器及配套服务。",
        purchaser=purchaser,
        budget=budget,
        deadline=deadline,
        raw_content_fingerprint=raw_fingerprint,
        notice_stable_fingerprint=notice_fingerprint,
        project_stable_fingerprint=project_fingerprint,
        fetched_at=fetched_at,
        evidence=evidence,
    )


class IncrementalDeliveryTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary_directory.name)
        self.database_path = self.root / "app.db"
        self.report_dir = self.root / "reports"
        self.task_id = f"task-{uuid4()}"
        self.patches = [
            patch.object(database_module, "DATA_DIR", self.root),
            patch.object(database_module, "DATABASE_PATH", self.database_path),
            patch.object(publisher_module, "REPORT_DIR", self.report_dir),
        ]
        for active_patch in self.patches:
            active_patch.start()

    def tearDown(self) -> None:
        for active_patch in reversed(self.patches):
            active_patch.stop()
        self.temporary_directory.cleanup()

    def run_workflow(self, adapters, *, run_id: str | None = None) -> dict:
        with patch.object(source_select, "SOURCE_ADAPTERS", adapters):
            return self.invoke_workflow(run_id=run_id)

    def invoke_workflow(self, *, run_id: str | None = None) -> dict:
        return asyncio.run(
            WORKFLOW.ainvoke(
                {
                    "task_id": self.task_id,
                    "run_id": run_id or str(uuid4()),
                    "query": "查询 2026-07-14 全国服务器采购公告",
                    "frequency": "daily",
                    "status": "running",
                    "steps": [],
                    "funnel": {},
                    "retry_count": 0,
                    "quality_passed": False,
                    "quality_issues": [],
                },
                config={"recursion_limit": 50},
            )
        )

    @staticmethod
    def successful_source(notices: list[TenderNotice], source_id: str = "public-a"):
        return SuccessfulSource(source_metadata(source_id, f"公开来源 {source_id}"), notices)

    def test_first_run_commits_a_full_snapshot_delivery(self) -> None:
        state = self.run_workflow([self.successful_source([make_notice()])])

        self.assertEqual(state["report"]["status"], "generated")
        self.assertEqual(state["report"]["delivery_type"], "full_snapshot")
        self.assertEqual(state["report"]["notice_count"], 1)
        self.assertEqual(len(list(self.report_dir.glob("*.docx"))), 1)

        repository = Repository()
        deliveries = repository.list_deliveries(self.task_id)
        self.assertEqual(len(deliveries), 1)
        self.assertEqual(deliveries[0]["status"], "generated")
        self.assertEqual(deliveries[0]["project_stable_fingerprints"], ["1" * 64])
        self.assertEqual(deliveries[0]["notice_stable_fingerprints"], ["2" * 64])
        snapshots = repository.load_project_snapshots(self.task_id)
        self.assertEqual(len(snapshots), 1)
        snapshot = next(iter(snapshots.values()))
        self.assertEqual(snapshot["version"], 1)
        self.assertTrue(snapshot["notices"][0]["evidence"])
        self.assertEqual(
            repository.load_watermarks(self.task_id)["public-a"]["source_id"],
            "public-a",
        )

    def test_identical_second_run_returns_no_change_without_a_new_file(self) -> None:
        adapter = self.successful_source([make_notice()])
        first = self.run_workflow([adapter])
        second = self.run_workflow([adapter])

        self.assertEqual(first["report"]["status"], "generated")
        self.assertEqual(second["report"]["status"], "no_change")
        self.assertEqual(second["report"]["notice_count"], 0)
        self.assertIsNone(second["report"]["filename"])
        self.assertIsNone(second["report"]["download_url"])
        self.assertIsNotNone(second["report"]["historical_report"])
        self.assertEqual(len(list(self.report_dir.glob("*.docx"))), 1)
        self.assertEqual(len(Repository().list_deliveries(self.task_id)), 1)
        snapshot = next(iter(Repository().load_project_snapshots(self.task_id).values()))
        self.assertEqual(snapshot["version"], 1)

    def test_a_new_project_is_the_only_notice_in_the_incremental_report(self) -> None:
        original = make_notice()
        added = make_notice(
            source_url="https://public-a.gov.cn/notices/002",
            project_fingerprint="4" * 64,
            notice_fingerprint="5" * 64,
            raw_fingerprint="6" * 64,
            title="新增存储设备采购公告",
        )
        self.run_workflow([self.successful_source([original])])
        state = self.run_workflow([self.successful_source([original, added])])

        self.assertEqual(state["report"]["status"], "generated")
        self.assertEqual(state["report"]["delivery_type"], "new_project")
        self.assertEqual(state["report"]["report_scope"], "incremental")
        self.assertEqual(state["report"]["notice_count"], 1)
        self.assertEqual([item["project_id"] for item in state["changes"]], ["project-4444444444444444"])
        report_path = self.report_dir / state["report"]["filename"]
        text = "\n".join(
            paragraph.text
            for table in Document(report_path).tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        )
        self.assertIn("新增存储设备采购公告", text)
        self.assertNotIn("某单位服务器采购公告", text)

    def test_material_fields_and_lifecycle_keep_before_after_values_and_evidence(self) -> None:
        original = make_notice()
        correction = make_notice(
            source_url="https://public-a.gov.cn/notices/001-correction",
            notice_fingerprint="7" * 64,
            raw_fingerprint="8" * 64,
            notice_type="correction",
            budget="1200000.00",
            deadline="2026-08-05T17:00:00+08:00",
            purchaser="某采购单位（新名称）",
            title="某单位服务器采购更正公告",
            published_at="2026-07-15T09:00:00+08:00",
            fetched_at="2026-07-15T10:00:00+08:00",
        )
        self.run_workflow([self.successful_source([original])])
        state = self.run_workflow([self.successful_source([original, correction])])

        self.assertEqual(state["report"]["delivery_type"], "material_change")
        self.assertEqual(len(state["changes"]), 1)
        field_changes = {
            item["field_path"]: item for item in state["changes"][0]["fields"]
        }
        self.assertEqual(
            set(field_changes),
            {"budget", "deadline", "purchaser", "notice_lifecycle"},
        )
        self.assertEqual(field_changes["budget"]["previous_value"], "1000000.00")
        self.assertEqual(field_changes["budget"]["current_value"], "1200000.00")
        self.assertEqual(field_changes["deadline"]["previous_value"], "2026-07-31T17:00:00+08:00")
        self.assertEqual(field_changes["deadline"]["current_value"], "2026-08-05T17:00:00+08:00")
        self.assertEqual(field_changes["purchaser"]["previous_value"], "某采购单位")
        self.assertEqual(field_changes["purchaser"]["current_value"], "某采购单位（新名称）")
        self.assertTrue(all(item["evidence_ids"] for item in field_changes.values()))
        self.assertTrue(all(item["evidence"] for item in field_changes.values()))
        snapshot = next(iter(Repository().load_project_snapshots(self.task_id).values()))
        self.assertEqual(snapshot["version"], 2)

    def test_format_fetch_time_and_mirror_source_changes_are_not_material(self) -> None:
        original = make_notice()
        mirror_refresh = make_notice(
            source_id="public-b",
            source_url="https://public-b.gov.cn/reposts/001",
            raw_fingerprint="9" * 64,
            budget="1000000.0",
            purchaser="某 采购单位",
            fetched_at="2026-07-16T11:30:00+08:00",
        )
        self.run_workflow([self.successful_source([original])])
        state = self.run_workflow(
            [self.successful_source([mirror_refresh], source_id="public-b")]
        )

        self.assertEqual(state["changes"], [])
        self.assertEqual(state["report"]["status"], "no_change")
        self.assertIsNone(state["report"]["filename"])
        self.assertEqual(len(list(self.report_dir.glob("*.docx"))), 1)
        self.assertEqual(len(Repository().list_deliveries(self.task_id)), 1)
        snapshot = next(iter(Repository().load_project_snapshots(self.task_id).values()))
        self.assertEqual(snapshot["version"], 1)

    def test_docx_failure_does_not_advance_snapshot_or_watermark_and_retry_sees_change(self) -> None:
        original = make_notice()
        changed = make_notice(
            raw_fingerprint="a" * 64,
            budget="1300000.00",
            fetched_at="2026-07-15T10:00:00+08:00",
        )
        self.run_workflow([self.successful_source([original])])
        repository = Repository()
        initial_snapshot = next(iter(repository.load_project_snapshots(self.task_id).values()))
        initial_watermark = repository.load_watermarks(self.task_id)["public-a"]

        with patch.object(
            publisher_module.DocxPublisher,
            "publish",
            side_effect=RuntimeError("injected DOCX failure"),
        ):
            failed = self.run_workflow([self.successful_source([changed])])

        self.assertEqual(failed["status"], "failed")
        self.assertEqual(failed["report"]["status"], "failed")
        self.assertEqual(failed["report"]["error_type"], "RuntimeError")
        after_failure = Repository()
        failed_snapshot = next(iter(after_failure.load_project_snapshots(self.task_id).values()))
        self.assertEqual(failed_snapshot["version"], initial_snapshot["version"])
        self.assertEqual(failed_snapshot["facts"], initial_snapshot["facts"])
        self.assertEqual(
            after_failure.load_watermarks(self.task_id)["public-a"],
            initial_watermark,
        )
        self.assertEqual(len(list(self.report_dir.glob("*.docx"))), 1)

        retried = self.run_workflow([self.successful_source([changed])])

        self.assertEqual(retried["status"], "completed")
        self.assertEqual(retried["report"]["delivery_type"], "material_change")
        self.assertEqual(retried["changes"][0]["fields"][0]["previous_value"], "1000000.00")
        self.assertEqual(retried["changes"][0]["fields"][0]["current_value"], "1300000.00")
        self.assertEqual(len(list(self.report_dir.glob("*.docx"))), 2)
        deliveries = Repository().list_deliveries(self.task_id)
        self.assertEqual(len(deliveries), 2)
        self.assertEqual([item["status"] for item in deliveries], ["generated", "generated"])
        final_snapshot = next(iter(Repository().load_project_snapshots(self.task_id).values()))
        self.assertEqual(final_snapshot["version"], 2)

    def test_only_successful_sources_advance_their_watermarks(self) -> None:
        adapters = [
            self.successful_source([make_notice()], source_id="public-a"),
            FailingSource(source_metadata("public-b", "公开来源 public-b")),
        ]

        state = self.run_workflow(adapters)

        self.assertEqual(state["status"], "completed")
        watermarks = Repository().load_watermarks(self.task_id)
        self.assertEqual(set(watermarks), {"public-a"})
        self.assertEqual(watermarks["public-a"]["record_count"], 1)
        self.assertEqual(state["report"]["failed_sources"][0]["source_id"], "public-b")

    def test_concurrent_runs_create_one_delivery_record_and_one_docx(self) -> None:
        adapters = [self.successful_source([make_notice()])]
        with patch.object(source_select, "SOURCE_ADAPTERS", adapters):
            with ThreadPoolExecutor(max_workers=2) as executor:
                states = list(
                    executor.map(
                        lambda run_id: self.invoke_workflow(run_id=run_id),
                        ("run-concurrent-a", "run-concurrent-b"),
                    )
                )

        self.assertTrue(all(state["status"] == "completed" for state in states))
        self.assertEqual(
            {state["report"]["delivery_fingerprint"] for state in states},
            {states[0]["report"]["delivery_fingerprint"]},
        )
        self.assertEqual(sum(state["report"]["reused_artifact"] for state in states), 1)
        self.assertEqual(len(Repository().list_deliveries(self.task_id)), 1)
        self.assertEqual(len(list(self.report_dir.glob("*.docx"))), 1)

    def test_prototype_database_migration_is_repeatable_and_preserves_rows(self) -> None:
        with database_module.connect() as connection:
            connection.executescript(
                """
                CREATE TABLE project_snapshots (
                    task_id TEXT NOT NULL,
                    project_id TEXT NOT NULL,
                    facts_json TEXT NOT NULL,
                    updated_at TEXT NOT NULL,
                    PRIMARY KEY (task_id, project_id)
                );
                INSERT INTO project_snapshots(task_id, project_id, facts_json, updated_at)
                VALUES ('legacy-task', 'legacy-project', '{"budget":"42"}', '2026-07-13T00:00:00+08:00');
                """
            )

        database_module.initialize_database()
        database_module.initialize_database()

        legacy = Repository().load_project_snapshots("legacy-task")
        self.assertIn("legacy-project", legacy)
        self.assertEqual(legacy["legacy-project"]["facts"]["budget"], "42")
        with database_module.connect() as connection:
            columns = {
                row["name"]
                for row in connection.execute("PRAGMA table_info(project_snapshots)")
            }
            tables = {
                row["name"]
                for row in connection.execute(
                    "SELECT name FROM sqlite_master WHERE type = 'table'"
                )
            }
        self.assertTrue(
            {"project_stable_fingerprint", "snapshot_fingerprint", "snapshot_json", "version"}.issubset(columns)
        )
        self.assertTrue({"deliveries", "source_watermarks", "notice_snapshots"}.issubset(tables))

    def test_stale_lock_and_staging_directory_do_not_block_a_run(self) -> None:
        notice = make_notice()
        project = {
            "project_id": "project-1111111111111111",
            "project_stable_fingerprint": "1" * 64,
            "documents": [{"notice": notice.model_dump(mode="json"), "is_mirror": False}],
        }
        snapshots = ChangeDetector().build_snapshots([project])
        changes = ChangeDetector().compare(snapshots, {})
        delivery_fingerprint = Publisher._delivery_fingerprint(
            {"task_id": self.task_id, "changes": changes},
            report_scope="full",
            report_projects=[project],
        )
        sample_name = publisher_module.build_report_filename(
            "查询 2026-07-14 全国服务器采购公告",
            datetime.now().astimezone(),
        )
        safe_query = sample_name.rsplit("_", maxsplit=1)[0]
        report_path = self.report_dir / f"{safe_query}_{self.task_id}_{delivery_fingerprint}.docx"
        self.report_dir.mkdir(parents=True)
        lock_path = report_path.with_suffix(".lock")
        lock_path.write_bytes(b"legacy")
        staging_dir = self.report_dir / ".staging" / "abandoned-run"
        staging_dir.mkdir(parents=True)
        (staging_dir / "partial.docx").write_bytes(b"partial")
        stale_time = time.time() - 3600
        os.utime(lock_path, (stale_time, stale_time))
        os.utime(staging_dir, (stale_time, stale_time))

        state = self.run_workflow([self.successful_source([notice])])

        self.assertEqual(state["status"], "completed")
        self.assertTrue(report_path.is_file())
        self.assertFalse(lock_path.exists())
        self.assertFalse(staging_dir.exists())

    def test_retry_fingerprint_ignores_refetched_evidence_metadata(self) -> None:
        original = make_notice()
        changed_once = make_notice(
            raw_fingerprint="c" * 64,
            budget="1700000.00",
            fetched_at="2026-07-15T10:00:00+08:00",
        )
        changed_refetched = make_notice(
            raw_fingerprint="d" * 64,
            budget="1700000.00",
            fetched_at="2026-07-16T10:00:00+08:00",
        )
        self.run_workflow([self.successful_source([original])])
        with patch.object(
            publisher_module.DocxPublisher,
            "publish",
            side_effect=RuntimeError("injected DOCX failure"),
        ):
            failed = self.run_workflow([self.successful_source([changed_once])])
        retried = self.run_workflow([self.successful_source([changed_refetched])])

        self.assertEqual(
            failed["report"]["delivery_fingerprint"],
            retried["report"]["delivery_fingerprint"],
        )
        self.assertEqual(len(Repository().list_deliveries(self.task_id)), 2)

    def test_new_only_batch_after_history_is_incremental_not_full(self) -> None:
        original = make_notice()
        added = make_notice(
            source_url="https://public-a.gov.cn/notices/new-only",
            project_fingerprint="e" * 64,
            notice_fingerprint="f" * 64,
            raw_fingerprint="a" * 64,
            title="历史任务中的后续新增项目",
        )
        self.run_workflow([self.successful_source([original])])
        state = self.run_workflow([self.successful_source([added])])

        self.assertEqual(state["report"]["delivery_type"], "new_project")
        self.assertEqual(state["report"]["report_scope"], "incremental")

    def test_watermark_never_regresses_on_an_older_successful_observation(self) -> None:
        newer = make_notice(
            raw_fingerprint="b" * 64,
            fetched_at="2026-07-16T10:00:00+08:00",
        )
        older = make_notice(
            raw_fingerprint="c" * 64,
            fetched_at="2026-07-15T10:00:00+08:00",
        )
        self.run_workflow([self.successful_source([newer])])
        self.run_workflow([self.successful_source([older])])

        watermark = Repository().load_watermarks(self.task_id)["public-a"]
        self.assertEqual(watermark["max_fetched_at"], "2026-07-16T10:00:00+08:00")

    def test_no_change_refresh_is_kept_as_a_notice_snapshot_version(self) -> None:
        original = make_notice()
        refreshed = make_notice(
            source_id="public-b",
            source_url="https://public-b.gov.cn/reposts/001",
            raw_fingerprint="d" * 64,
            fetched_at="2026-07-16T10:00:00+08:00",
        )
        self.run_workflow([self.successful_source([original])])
        state = self.run_workflow(
            [self.successful_source([refreshed], source_id="public-b")]
        )

        self.assertEqual(state["report"]["status"], "no_change")
        with database_module.connect() as connection:
            count = connection.execute(
                "SELECT COUNT(*) FROM notice_snapshots WHERE task_id = ?",
                (self.task_id,),
            ).fetchone()[0]
        self.assertEqual(count, 2)

    def test_stale_pending_delivery_can_be_reclaimed(self) -> None:
        repository = Repository()
        fingerprint = "f" * 64
        acquired, _ = repository.reserve_delivery(
            task_id=self.task_id,
            run_id="abandoned-run",
            delivery_type="full_snapshot",
            delivery_fingerprint=fingerprint,
            changes=[],
        )
        self.assertTrue(acquired)
        with database_module.connect() as connection:
            connection.execute(
                "UPDATE deliveries SET updated_at = ? WHERE delivery_fingerprint = ?",
                ("2026-07-14T00:00:00+08:00", fingerprint),
            )

        reclaimed, delivery = repository.reserve_delivery(
            task_id=self.task_id,
            run_id="recovery-run",
            delivery_type="full_snapshot",
            delivery_fingerprint=fingerprint,
            changes=[],
        )

        self.assertTrue(reclaimed)
        self.assertEqual(delivery["run_id"], "recovery-run")


if __name__ == "__main__":
    unittest.main()
